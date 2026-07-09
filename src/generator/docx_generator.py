"""
DOCX Generator - Füllt Spesenabrechnung-Vorlage mit Daten
Verbesserte Version mit korrekter Checkbox-Formatierung
"""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.logger import setup_logger
from utils.match_utils import parse_anpfiff, generate_filename_from_match, sanitize_team_name, extract_iso_date_from_anpfiff
from generator.spesen_calculator import calculate_spesen, format_spesen

logger = setup_logger("docx_generator")

# Kilometerpauschale laut Formular
KM_SATZ_EURO = 0.30

# Checkbox-Zeichen (beide Unicode, beide mit Segoe UI Symbol darstellbar)
CHECKBOX_CHECKED = '☒'    # U+2612 - Ballot Box with X
CHECKBOX_UNCHECKED = '☐'  # U+2610 - Ballot Box (leer)
CHECKBOX_FONT = 'Segoe UI Symbol'


def format_name_nachname_vorname(name: str) -> str:
    """
    Wandelt einen gescrapten Namen im Format "Vorname Nachname" in das im
    Formular erwartete Format "Nachname, Vorname" um.
    Letztes Wort gilt als Nachname, der Rest als Vorname.
    """
    parts = name.strip().split()
    if len(parts) < 2:
        return name
    return f"{parts[-1]}, {' '.join(parts[:-1])}"


class SpesenGenerator:
    """Generiert ausgefüllte Spesenabrechnung-Dokumente"""

    def __init__(self, template_path: str, output_dir: Path):
        """
        Initialisiert den Generator.

        Args:
            template_path: Pfad zur DOCX-Vorlage
            output_dir: Verzeichnis für generierte Dokumente (Path-Objekt vom SessionManager)
        """
        self.template_path = Path(template_path)
        self.output_dir = Path(output_dir)

        self.output_dir.mkdir(exist_ok=True, parents=True)

        if not self.template_path.exists():
            raise FileNotFoundError(f"Vorlage nicht gefunden: {self.template_path}")

        logger.info(f"Generator initialisiert mit Vorlage: {self.template_path}")
        logger.info(f"Output-Verzeichnis: {self.output_dir}")

    def _determine_checkboxes(self, match_data: dict) -> dict:
        """Bestimmt welche Checkboxen aktiviert werden müssen."""
        spiel_info = match_data.get('spiel_info', {})

        checkbox_keys = [
            'CHECKBOX_PUNKTSPIEL', 'CHECKBOX_POKALSPIEL', 'CHECKBOX_ENTSCHEIDUNG',
            'CHECKBOX_FREUNDSCHAFT', 'CHECKBOX_MAENNER', 'CHECKBOX_FRAUEN',
            'CHECKBOX_MAEDCHEN', 'CHECKBOX_ALTE_HERREN', 'CHECKBOX_SONSTIGE',
            'CHECKBOX_A_JUN', 'CHECKBOX_B_JUN', 'CHECKBOX_C_JUN',
            'CHECKBOX_D_JUN', 'CHECKBOX_E_JUN', 'CHECKBOX_F_JUN',
        ]
        checkboxes = {key: False for key in checkbox_keys}

        # Spielklasse prüfen
        spielklasse = spiel_info.get('spielklasse', '').lower()
        if 'pokal' in spielklasse:
            checkboxes['CHECKBOX_POKALSPIEL'] = True
        elif 'freundschaft' in spielklasse:
            checkboxes['CHECKBOX_FREUNDSCHAFT'] = True
        else:
            checkboxes['CHECKBOX_PUNKTSPIEL'] = True

        # Mannschaftsart prüfen
        mannschaftsart = spiel_info.get('mannschaftsart', '').lower()
        if 'herren' in mannschaftsart or 'männer' in mannschaftsart:
            if 'alte' in mannschaftsart:
                checkboxes['CHECKBOX_ALTE_HERREN'] = True
            else:
                checkboxes['CHECKBOX_MAENNER'] = True
        elif 'frauen' in mannschaftsart:
            checkboxes['CHECKBOX_FRAUEN'] = True
        elif 'mädchen' in mannschaftsart:
            checkboxes['CHECKBOX_MAEDCHEN'] = True
        elif 'a-junioren' in mannschaftsart:
            checkboxes['CHECKBOX_A_JUN'] = True
        elif 'b-junioren' in mannschaftsart:
            checkboxes['CHECKBOX_B_JUN'] = True
        elif 'c-junioren' in mannschaftsart:
            checkboxes['CHECKBOX_C_JUN'] = True
        elif 'd-junioren' in mannschaftsart:
            checkboxes['CHECKBOX_D_JUN'] = True
        elif 'e-junioren' in mannschaftsart:
            checkboxes['CHECKBOX_E_JUN'] = True
        elif 'f-junioren' in mannschaftsart:
            checkboxes['CHECKBOX_F_JUN'] = True
        else:
            checkboxes['CHECKBOX_SONSTIGE'] = True

        return checkboxes

    def _get_referee_by_role(self, referees: list, role: str) -> dict:
        """Findet Schiedsrichter nach Rolle."""
        for ref in referees:
            if ref.get('rolle') == role:
                return ref
        return {}

    def _set_run_font(self, run, font_name: str):
        """Setzt die Schriftart eines Runs vollständig."""
        run.font.name = font_name
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)

    def _replace_in_paragraph(self, paragraph, checkbox_states: dict, text_replacements: dict):
        """
        Ersetzt Platzhalter in einem Paragraph.
        Behandelt auch Platzhalter die über mehrere Runs verteilt sind.
        """
        # Alle Ersetzungen sammeln
        all_replacements = {}

        for key, is_checked in checkbox_states.items():
            placeholder = f"{{{{{key}}}}}"
            char = CHECKBOX_CHECKED if is_checked else CHECKBOX_UNCHECKED
            all_replacements[placeholder] = (char, True)  # True = ist Checkbox

        for key, value in text_replacements.items():
            placeholder = f"{{{{{key}}}}}"
            all_replacements[placeholder] = (str(value), False)  # False = ist Text

        # Für jeden Platzhalter prüfen und ersetzen
        for placeholder, (replacement, is_checkbox) in all_replacements.items():
            self._replace_placeholder_in_paragraph(paragraph, placeholder, replacement, is_checkbox)

    def _replace_placeholder_in_paragraph(self, paragraph, placeholder: str, replacement: str, is_checkbox: bool):
        """
        Ersetzt einen Platzhalter im Paragraph - auch wenn er über mehrere Runs verteilt ist.
        """
        runs = paragraph.runs
        if not runs:
            return

        # Gesamttext und Run-Grenzen ermitteln
        full_text = ''.join(run.text for run in runs)

        if placeholder not in full_text:
            return

        # Finde Position des Platzhalters
        start_idx = full_text.find(placeholder)
        end_idx = start_idx + len(placeholder)

        # Finde welche Runs betroffen sind
        char_count = 0
        start_run_idx = None
        end_run_idx = None
        start_char_in_run = None
        end_char_in_run = None

        for i, run in enumerate(runs):
            run_start = char_count
            run_end = char_count + len(run.text)

            # Start des Platzhalters in diesem Run?
            if start_run_idx is None and run_start <= start_idx < run_end:
                start_run_idx = i
                start_char_in_run = start_idx - run_start

            # Ende des Platzhalters in diesem Run?
            if end_run_idx is None and run_start < end_idx <= run_end:
                end_run_idx = i
                end_char_in_run = end_idx - run_start

            char_count = run_end

        if start_run_idx is None or end_run_idx is None:
            return

        # Fall 1: Platzhalter komplett in einem Run
        if start_run_idx == end_run_idx:
            run = runs[start_run_idx]
            run.text = run.text[:start_char_in_run] + replacement + run.text[end_char_in_run:]
            if is_checkbox:
                self._set_run_font(run, CHECKBOX_FONT)

        # Fall 2: Platzhalter über mehrere Runs verteilt
        else:
            # Ersten Run: Text vor Platzhalter + Ersetzung
            first_run = runs[start_run_idx]
            first_run.text = first_run.text[:start_char_in_run] + replacement
            if is_checkbox:
                self._set_run_font(first_run, CHECKBOX_FONT)

            # Mittlere Runs: komplett leeren
            for i in range(start_run_idx + 1, end_run_idx):
                runs[i].text = ''

            # Letzten Run: nur Text nach Platzhalter behalten
            last_run = runs[end_run_idx]
            last_run.text = last_run.text[end_char_in_run:]

    def _replace_placeholders(self, doc: Document, checkbox_states: dict, text_replacements: dict):
        """Ersetzt alle Platzhalter im Dokument."""
        # In Paragraphs ersetzen
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, checkbox_states, text_replacements)

        # In Tabellen ersetzen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, checkbox_states, text_replacements)

    def _calculate_spesen_for_match(self, match_data: dict, is_punktspiel: bool) -> tuple:
        """Berechnet Spesen für ein Spiel."""
        if not is_punktspiel:
            return ("", "")

        spiel_info = match_data.get('spiel_info', {})
        spielklasse = spiel_info.get('spielklasse', '')
        mannschaftsart = spiel_info.get('mannschaftsart', '')

        sr_spesen, sra_spesen = calculate_spesen(spielklasse, mannschaftsart)
        sr_spesen_str = format_spesen(sr_spesen)
        sra_spesen_str = format_spesen(sra_spesen)

        if sr_spesen_str:
            logger.info(f"Spesen berechnet: SR={sr_spesen_str}, SRA={sra_spesen_str or '-'}")

        return (sr_spesen_str, sra_spesen_str)

    def _build_expense_replacements(self, expenses: dict, is_punktspiel: bool,
                                    match_data: dict, sra1_active: bool, sra2_active: bool) -> dict:
        """
        Baut die Platzhalter fuer Fahrtkosten (km x 0,30 EUR), OeVM und die
        Summen-Zeile (Spesen + Fahrtkosten + OeVM) pro Person.
        """
        expenses = expenses or {}

        # Numerische Spesen fuer die Summenberechnung
        sr_spesen_num, sra_spesen_num = (None, None)
        if is_punktspiel:
            spiel_info = match_data.get('spiel_info', {})
            sr_spesen_num, sra_spesen_num = calculate_spesen(
                spiel_info.get('spielklasse', ''), spiel_info.get('mannschaftsart', '')
            )

        replacements = {}
        roles = [
            ('SR', 'sr', sr_spesen_num, True),
            ('SRA1', 'sra1', sra_spesen_num, sra1_active),
            ('SRA2', 'sra2', sra_spesen_num, sra2_active),
        ]

        # Fuer die Gesamtsumme: alle angesetzten Personen muessen erfasst sein
        # (eine explizite 0 zaehlt als erfasst, ein leeres Feld nicht)
        alle_aktiven_erfasst = True
        einzelsummen = []

        for prefix, key, spesen_num, active in roles:
            km = expenses.get(f'{key}_km')
            oevm = expenses.get(f'{key}_oevm')

            # 0 ist ein gueltiger Wert (bewusst erfasst), None = nicht eingetragen
            km_kosten = round(km * KM_SATZ_EURO, 2) if km is not None else None
            oevm_betrag = oevm if oevm is not None else None

            # Anzeige: 0-Betraege bleiben im Formular leer
            replacements[f'{prefix}_Kilometer'] = format_spesen(km_kosten) if km_kosten else ''
            replacements[f'{prefix}_OEVM'] = format_spesen(oevm_betrag) if oevm_betrag else ''

            # Summe nur, wenn die Person angesetzt ist und mindestens eine Komponente existiert
            komponenten = [v for v in (spesen_num if active else None, km_kosten, oevm_betrag) if v is not None]
            summe = round(sum(komponenten), 2) if komponenten and active else None
            replacements[f'{prefix}_Summe'] = format_spesen(summe) if summe else ''

            if active:
                erfasst = km is not None or oevm is not None
                if not erfasst:
                    alle_aktiven_erfasst = False
                elif summe is not None:
                    einzelsummen.append(summe)

        # Gesamtsumme ueber alle angesetzten Personen
        if alle_aktiven_erfasst and einzelsummen:
            replacements['Summe'] = format_spesen(round(sum(einzelsummen), 2))
        else:
            replacements['Summe'] = ''

        return replacements

    def generate_document(self, match_data: dict, output_filename: str = None, expenses: dict = None) -> Path:
        """
        Generiert ein ausgefülltes Dokument für ein Spiel.

        Args:
            expenses: Optionale Fahrtkosten/OeVM-Werte
                      (Keys: sr_km, sr_oevm, sra1_km, sra1_oevm, sra2_km, sra2_oevm)
        """
        spiel_info = match_data.get('spiel_info', {})
        schiedsrichter = match_data.get('schiedsrichter', [])
        spielstaette = match_data.get('spielstaette', {})

        doc = Document(self.template_path)
        logger.debug(f"Vorlage geladen für: {spiel_info.get('heim_team', '')} vs {spiel_info.get('gast_team', '')}")

        datum, anstoss = parse_anpfiff(spiel_info.get('anpfiff', ''))
        checkboxes = self._determine_checkboxes(match_data)

        is_punktspiel = checkboxes['CHECKBOX_PUNKTSPIEL']
        sr_spesen_str, sra_spesen_str = self._calculate_spesen_for_match(match_data, is_punktspiel)

        sr = self._get_referee_by_role(schiedsrichter, 'SR')
        sra1 = self._get_referee_by_role(schiedsrichter, 'SRA 1')
        sra2 = self._get_referee_by_role(schiedsrichter, 'SRA 2')

        spielort_name = spielstaette.get('name', '')
        spielort_adresse = spielstaette.get('adresse', '')
        spielort_komplett = f"{spielort_name}\n{spielort_adresse}"

        sra1_spesen = sra_spesen_str if sra1.get('name') else ''
        sra2_spesen = sra_spesen_str if sra2.get('name') else ''

        expense_replacements = self._build_expense_replacements(
            expenses, is_punktspiel, match_data,
            sra1_active=bool(sra1.get('name')), sra2_active=bool(sra2.get('name'))
        )

        text_replacements = {
            **expense_replacements,
            'HEIM_TEAM': spiel_info.get('heim_team', ''),
            'GAST_TEAM': spiel_info.get('gast_team', ''),
            'SPIELKLASSE': spiel_info.get('spielklasse', ''),
            'SPIELNUMMER': '',
            'DATUM': datum,
            'ANSTOSS': anstoss,
            'SPIELORT': spielort_komplett,
            'SR_NAME': format_name_nachname_vorname(sr.get('name', '')),
            'SR_STRASSE': sr.get('strasse', ''),
            'SR_PLZ_ORT': sr.get('plz_ort', ''),
            'SRA1_NAME': format_name_nachname_vorname(sra1.get('name', '')),
            'SRA1_STRASSE': sra1.get('strasse', ''),
            'SRA1_PLZ_ORT': sra1.get('plz_ort', ''),
            'SRA2_NAME': format_name_nachname_vorname(sra2.get('name', '')),
            'SRA2_STRASSE': sra2.get('strasse', ''),
            'SRA2_PLZ_ORT': sra2.get('plz_ort', ''),
            'SR_Spesen': sr_spesen_str,
            'SR1_Spesen': sra1_spesen,
            'SR2_Spesen': sra2_spesen,
        }

        self._replace_placeholders(doc, checkboxes, text_replacements)

        if not output_filename:
            output_filename = generate_filename_from_match(match_data)

        output_path = self.output_dir / output_filename
        doc.save(output_path)

        logger.info(f"Dokument erstellt: {output_path}")
        return output_path

    def generate_all_documents(self, matches_data: list, expenses_map: dict = None) -> list:
        """
        Generiert Dokumente für alle Spiele.

        Args:
            expenses_map: Optionales Dict {(heim, gast, datum_iso): expenses}
                          mit gespeicherten Fahrtkosten/OeVM pro Spiel.
        """
        logger.info(f"Generiere {len(matches_data)} Dokumente...")
        expenses_map = expenses_map or {}

        generated_files = []
        for i, match_data in enumerate(matches_data, 1):
            try:
                spiel_info = match_data.get('spiel_info', {})
                heim = spiel_info.get('heim_team', 'Unbekannt')
                gast = spiel_info.get('gast_team', 'Unbekannt')
                datum_iso = extract_iso_date_from_anpfiff(spiel_info.get('anpfiff', ''))
                expenses = expenses_map.get((heim, gast, datum_iso))

                logger.info(f"[{i}/{len(matches_data)}] Verarbeite: {heim} vs {gast}")
                output_path = self.generate_document(match_data, expenses=expenses)
                generated_files.append(output_path)
                logger.info(f"[{i}/{len(matches_data)}] ✓ Erstellt")
            except Exception as e:
                logger.error(f"[{i}/{len(matches_data)}] ✗ Fehler: {e}")
                continue

        logger.info(f"Erfolgreich {len(generated_files)}/{len(matches_data)} Dokumente generiert")
        return generated_files